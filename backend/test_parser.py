import os
import urllib.request
import docx
from parser import extract_text

def create_sample_docx(file_path):
    print(f"Creating sample DOCX at: {file_path}")
    doc = docx.Document()
    doc.add_heading("Jane Doe Resume", 0)
    doc.add_paragraph("Email: jane.doe@example.com")
    doc.add_paragraph("Experience: Senior Software Engineer at Tech Corp")
    doc.add_paragraph("Skills: Python, FastAPI, React")
    
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Project'
    hdr_cells[1].text = 'Description'
    row_cells = table.rows[1].cells
    row_cells[0].text = 'AI Resume Analyzer'
    row_cells[1].text = 'Extracted text using python-docx.'
    
    doc.save(file_path)

def download_sample_pdf(file_path):
    url = "https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf"
    print(f"Downloading sample PDF from {url} to {file_path}")
    try:
        urllib.request.urlretrieve(url, file_path)
    except Exception as e:
        print(f"Warning: PDF download failed ({e}). Attempting to create a local text-based PDF or bypass.")
        raise e

def main():
    docx_path = "test_resume.docx"
    pdf_path = "test_resume.pdf"
    
    success = True
    
    # 1. Test DOCX Parsing
    try:
        create_sample_docx(docx_path)
        docx_text = extract_text(docx_path)
        print("\n--- Extracted DOCX Content ---")
        print(docx_text)
        print("------------------------------")
        assert "Jane Doe Resume" in docx_text
        assert "FastAPI" in docx_text
        assert "AI Resume Analyzer" in docx_text
        print("✓ DOCX parser test passed!")
    except Exception as e:
        print(f"✗ DOCX parser test failed: {e}")
        success = False
        
    # 2. Test PDF Parsing
    try:
        download_sample_pdf(pdf_path)
        pdf_text = extract_text(pdf_path)
        print("\n--- Extracted PDF Content ---")
        print(pdf_text)
        print("-----------------------------")
        assert "Dummy PDF file" in pdf_text or len(pdf_text.strip()) > 0
        print("✓ PDF parser test passed!")
    except Exception as e:
        print(f"✗ PDF parser test failed: {e}")
        success = False

    # 3. Test Unsupported Exception
    try:
        print("\nTesting unsupported extension exception...")
        # Create test.txt file so it exists
        with open("test.txt", "w") as f:
            f.write("unsupported content")
        extract_text("test.txt")
        print("✗ Unsupported test failed (no exception raised)")
        success = False
    except ValueError as e:
        print(f"✓ Unsupported test passed: {str(e)}")
    except Exception as e:
        print(f"✗ Unsupported test raised wrong exception: {type(e)}")
        success = False
    finally:
        if os.path.exists("test.txt"):
            os.remove("test.txt")

    # Clean up test files
    for path in [docx_path, pdf_path]:
        if os.path.exists(path):
            os.remove(path)
            print(f"Cleaned up {path}")

    if success:
        print("\n🎉 ALL PARSER TESTS PASSED SUCCESSFULLY!")
        exit(0)
    else:
        print("\n❌ SOME PARSER TESTS FAILED.")
        exit(1)

if __name__ == "__main__":
    main()
