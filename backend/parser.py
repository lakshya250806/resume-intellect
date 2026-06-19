import os
import re
import pdfplumber
import docx

SKILL_KEYWORDS = [
    "Python", "Java", "C++", "JavaScript", "TypeScript", "React", "Next.js", 
    "Node.js", "Express", "FastAPI", "Flask", "Spring Boot", "TensorFlow", 
    "PyTorch", "Docker", "Kubernetes", "Git", "GitHub", "SQL", "MongoDB", 
    "PostgreSQL", "Redis", "HTML", "CSS", "Tailwind", "Bootstrap", "Pandas", 
    "NumPy", "Scikit-learn", "OpenCV", "Gemini", "OpenAI", "AWS", "Azure", "Linux"
]

def extract_text(file_path: str) -> str:
    """
    Extracts raw plain text from a PDF or DOCX file.
    
    :param file_path: The absolute or relative path to the file.
    :return: The extracted plain text as a string.
    :raises FileNotFoundError: If the file does not exist.
    :raises ValueError: If the file format is unsupported.
    :raises RuntimeError: If text extraction fails due to file corruption or parsing issues.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at: {file_path}")

    _, ext = os.path.splitext(file_path.lower())
    
    if ext == '.pdf':
        return _extract_pdf_text(file_path)
    elif ext == '.docx':
        return _extract_docx_text(file_path)
    else:
        raise ValueError(
            f"Unsupported file format '{ext}'. Only .pdf and .docx files are supported."
        )

def _extract_pdf_text(file_path: str) -> str:
    """
    Helper function to extract text from PDF files using pdfplumber.
    """
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            if not pdf.pages:
                raise ValueError("PDF file contains no pages.")
                
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text.strip())
        
        if not text_content:
            raise ValueError("No text could be extracted from the PDF pages (it might be scanned/image-only).")
            
        return "\n\n".join(text_content)
    except Exception as e:
        if isinstance(e, ValueError):
            raise
        raise RuntimeError(f"Failed to extract text from PDF file: {str(e)}") from e

def _extract_docx_text(file_path: str) -> str:
    """
    Helper function to extract text from Word DOCX files using python-docx.
    """
    try:
        doc = docx.Document(file_path)
        # Extract paragraph texts, ignoring empty paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Extract table cells to capture all text
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    # Remove duplicate adjacent cell texts (often happens with merged cells)
                    clean_row_text = []
                    for val in row_text:
                        if not clean_row_text or clean_row_text[-1] != val:
                            clean_row_text.append(val)
                    table_texts.append(" | ".join(clean_row_text))

        combined_text_list = paragraphs
        if table_texts:
            combined_text_list.append("\n--- Table Content ---\n" + "\n".join(table_texts))
            
        combined_text = "\n\n".join(combined_text_list)
        if not combined_text.strip():
            raise ValueError("DOCX file contains no text paragraphs or tables.")
            
        return combined_text
    except Exception as e:
        if isinstance(e, ValueError):
            raise
        raise RuntimeError(f"Failed to extract text from DOCX file: {str(e)}") from e

def infer_name(text: str) -> str:
    """
    Infors candidate's name from the first few lines of the resume.
    """
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    
    # Exclude lines that contain noise or sections headers
    exclusion_keywords = {
        "email", "phone", "skills", "experience", "education", "projects", 
        "summary", "objective", "profile", "contact", "address", "github", 
        "linkedin", "portfolio", "http", "www", "resume", "curriculum", "vitae",
        "about", "work", "history", "career"
    }
    
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    phone_pattern = r'\+?\d[\d\s\(\)-]{7,}\d'
    
    candidate_lines = []
    for line in lines[:10]:  # Look in the top 10 lines
        lower_line = line.lower()
        
        # Skip lines containing email, phone or common URL domains
        if re.search(email_pattern, line) or re.search(phone_pattern, line):
            continue
        if any(domain in lower_line for domain in ["github.com", "linkedin.com", "http", "www"]):
            continue
            
        # Skip if contains any exclusion keyword
        if any(kw in lower_line for kw in exclusion_keywords):
            continue
            
        # Clean lines of bullets or numbering
        cleaned = re.sub(r'^[\s\-\*•o\d\.\)]+', '', line).strip()
        if not cleaned:
            continue
            
        candidate_lines.append(cleaned)
        
    # Match standard capitalized name pattern (2 to 4 words)
    name_like_pattern = r'^[A-Z][a-zA-Z]*(?:\s+[A-Z][a-zA-Z]*){1,3}$'
    for candidate in candidate_lines:
        if re.match(name_like_pattern, candidate):
            return candidate
            
    # Fallback to the first candidate line that isn't excluded
    if candidate_lines:
        return candidate_lines[0]
        
    # Final fallback
    if lines:
        return lines[0]
    return ""

def extract_email(text: str) -> str:
    """
    Extracts the first email address found in the text.
    """
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    match = re.search(email_pattern, text)
    return match.group(0) if match else ""

def extract_phone(text: str) -> str:
    """
    Extracts the first phone number matching typical formats (Indian and International).
    """
    phone_patterns = [
        # Indian pattern split with dash/space: +91-98765-43210
        r'(?<!\w)(?:\+?91|0)?[-\s]?[6789]\d{4}[-\s]?\d{5}\b',
        # Indian pattern consecutive: +91 9876543210
        r'(?<!\w)(?:\+?91|0)?[-\s]?[6789]\d{9}\b',
        # General international formats like +1-555-123-4567, (555) 123-4567
        r'(?<!\w)(?:\+?\d{1,3}[-\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b',
        # Plain 10 digit number
        r'\b\d{10}\b'
    ]
    
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Clean spaces/dashes and return first match
            return re.sub(r'\s+', ' ', matches[0].strip())
    return ""

def extract_skills(text: str) -> list:
    """
    Matches predefined technical skills inside the raw text.
    """
    detected = []
    text_lower = text.lower()
    
    for skill in SKILL_KEYWORDS:
        escaped = re.escape(skill.lower())
        
        # Word boundary logic: use standard \b for alphanumeric, lookarounds for special chars like C++
        start_boundary = r'\b' if escaped[0].isalnum() else r'(?<!\w)'
        end_boundary = r'\b' if escaped[-1].isalnum() else r'(?!\w)'
        
        pattern = f"{start_boundary}{escaped}{end_boundary}"
        
        if re.search(pattern, text_lower):
            detected.append(skill)
            
    return sorted(list(set(detected)))

def extract_section_lines(text: str, regex_pattern: str) -> list:
    """
    Extracts and cleans lines matching a regex pattern.
    """
    matched_lines = []
    compiled = re.compile(regex_pattern, re.IGNORECASE)
    
    # Common section header titles to skip
    header_titles = {
        "education", "experience", "projects", "work experience", 
        "academic projects", "personal projects", "key projects", 
        "professional experience", "employment history", "employment", 
        "skills", "technical skills", "history", "career"
    }
    
    for line in text.split('\n'):
        cleaned = line.strip()
        if not cleaned:
            continue
            
        # Skip exact section header lines
        if cleaned.lower() in header_titles:
            continue
            
        if compiled.search(cleaned):
            # Strip list bullets, numbering from the start of the line
            cleaned = re.sub(r'^[\s\-\*•o\d\.\)]+', '', cleaned).strip()
            if cleaned and cleaned not in matched_lines:
                matched_lines.append(cleaned)
                
    return matched_lines

def extract_resume_info(text: str) -> dict:
    """
    Extracts structured summary information from raw resume text.
    
    :param text: Raw text extracted from the document.
    :return: Dict containing extracted profile components.
    """
    if not text:
        return {
            "name": "",
            "email": "",
            "phone": "",
            "skills": [],
            "education": [],
            "experience": [],
            "projects": []
        }
        
    education_pattern = r'\b(?:b\.tech|bachelor|master|university|institute|college|cgpa|diploma|ph\.d|school|degree)s?\b'
    experience_pattern = r'\b(?:engineer|developer|intern|analyst|experience|worked|company|employment|job|position|role)s?\b'
    projects_pattern = r'\b(?:project|built|developed|created|designed|implemented|github)s?\b'
    
    return {
        "name": infer_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills": extract_skills(text),
        "education": extract_section_lines(text, education_pattern),
        "experience": extract_section_lines(text, experience_pattern),
        "projects": extract_section_lines(text, projects_pattern)
    }

class ResumeParser:
    """
    A wrapper class to keep main.py compatibility.
    Accepts binary content and a filename, writes it to a temporary file,
    and extracts text using the core extract_text function.
    """
    def __init__(self):
        pass

    def parse(self, file_content: bytes, filename: str) -> dict:
        """
        Parses binary file contents by writing to a temporary file.
        
        :param file_content: Binary content of the file.
        :param filename: Name of the file with extension.
        :return: Dict with file metadata and extracted raw text.
        """
        import tempfile
        
        suffix = os.path.splitext(filename.lower())[1]
        if suffix not in ['.pdf', '.docx']:
            # Fallback to write temp file with same suffix to raise ValueError in extract_text
            suffix = '.unsupported'
            
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(file_content)
            temp_path = temp_file.name
            
        try:
            text = extract_text(temp_path)
            return {
                "filename": filename,
                "file_type": os.path.splitext(filename)[1].lstrip('.'),
                "file_size_bytes": len(file_content),
                "extracted_text": text
            }
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
